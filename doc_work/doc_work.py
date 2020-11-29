import os
import docx
import pdf2docx


def preprocess_text(text, remove_wrap=True, clear_paragraph=True):
    if remove_wrap:
        text = text.replace('-', '')
    if clear_paragraph:
        text = text.replace(' \n', '\n').replace('\n', ' ').replace('\t', '\n')
    return text


class Document:
    """
    Class 'Document' gets data_file_path
    Method 'parse' returns text from 'data_file_path' file
    """
    def __init__(self, data_file_path, text_params={'start': '', 'stop': '', 'fio': ''}):
        self.data_file_path = data_file_path
        self.filetype = data_file_path.rsplit('.')[-1]
        self.text_params = text_params
        self.text = ''
        self.docx = docx.Document()
        self.paragraphs = None
        self.start_target_paragraph = None
        self.stop_target_paragraph = None
        self.fio = None
        self.f = None

    def pdf2docx(self):
        pdf2docx.parse(self.data_file_path, self.data_file_path + '.docx')
        self.data_file_path = self.data_file_path + '.docx'
        self.filetype = 'docx'
        return None

    def txt2docx(self):
        for paragraph in self.paragraphs:
            self.docx.add_paragraph(paragraph)
        self.filetype = 'docx'
        return None

    def docx_parser(self):
        doc = docx.Document(self.data_file_path)
        self.docx = doc
        self.paragraphs = [paragraph.text for paragraph in doc.paragraphs]

    def txt_parser(self):
        text = open(self.data_file_path, 'r', encoding='utf-8').read()
        self.paragraphs = text.split('\n')

    def parse(self, void=False):
        # process doc
        if self.filetype == 'docx':
            self.docx_parser()
        elif self.filetype == 'pdf':
            self.pdf2docx()
            self.docx_parser()
        elif self.filetype == 'txt':
            self.txt_parser()
            self.txt2docx()
        # found target text (match start and stop paragraphs)
        for i, paragraph in enumerate(self.paragraphs):
            if self.start_target_paragraph is None:
                if paragraph.find(self.text_params['start']) != -1:
                    self.start_target_paragraph = i + 1
            if self.stop_target_paragraph is None:
                if paragraph.rfind(self.text_params['stop']) != -1:
                    self.stop_target_paragraph = i
            if self.fio is None:
                fio_pos_start = paragraph.find(self.text_params['fio'])
                if fio_pos_start != -1:
                    fio_pos_stop = paragraph.find('\n', fio_pos_start)
                    self.fio = paragraph[fio_pos_start: fio_pos_stop]
        # extract text
        self.paragraphs = self.paragraphs[self.start_target_paragraph: self.stop_target_paragraph]
        remove_wrap = self.filetype == 'pdf'
        self.paragraphs = list(map(lambda x: preprocess_text(x, remove_wrap), self.paragraphs))
        self.text = '\n'.join(self.paragraphs)
        if void:
            return None
        else:
            return self.text

    def find_fio(self, only_family_name=True):
        if self.filetype == 'docx':
            table_text = []
            for t in self.docx.tables:
                for r in t.rows:
                    for c in r.cells:
                        if self.fio is None:
                            fio_pos_start = c.text.find(self.text_params['fio'])
                            if fio_pos_start != -1:
                                self.fio = c.text[fio_pos_start + len(self.text_params['fio']):].strip()
                                self.f = self.fio.split()[0]
                        else:
                            break
        elif self.filetype == 'txt':
            fio_pos_start = self.text.find(self.text_params['fio'])
            if fio_pos_start != -1:
                self.fio = self.text[fio_pos_start + len(self.text_params['fio']):].strip()
                self.f = self.fio.split()[0]

        if only_family_name:
            return self.f
        else:
            return self.fio

    def change_text(self, new_text, highlight_text=True):
        if self.filetype == 'docx':
            new_text = new_text.split('\n')
            if self.start_target_paragraph is None and self.stop_target_paragraph is None:
                for i, text in zip(range(len(self.docx.paragraphs)), new_text):
                    self.docx.paragraphs[i].text = text
            else:
                for i, text in zip(range(self.start_target_paragraph, self.stop_target_paragraph), new_text):
                    if highlight_text:
                        t = ''
                        for w_old, w_new in zip(self.docx.paragraphs[i].text.split(), text.split()):
                            if w_old != w_new:
                                t = t + ' ' + str.upper(w_new)
                            else:
                                t = t + ' ' + w_new
                        self.docx.paragraphs[i].text = t
                    else:
                        self.docx.paragraphs[i].text = text
                    # for j in range(len(self.docx.paragraphs[i].runs)):
                    #     self.docx.paragraphs[i].runs[j].font.color.rgb = docx.shared.RGBColor(0x42, 0x24, 0xE9)
        elif self.filetype == 'txt':
            self.text = self.text[:self.start_target_paragraph] + new_text + self.text[:self.stop_target_paragraph]
        return None

    def save(self, new_name):
        self.docx.save(new_name)
        return None


if __name__ == '__main__':
    text_params = {'start': '307 УК РФ предупрежден',
                   'stop': 'Перед началом, в ходе либо по окончании',
                   'fio': 'Фамилия, имя, отчество'
                   }
    cwd = ['data']
    file_path = os.path.join(*cwd, 'input.docx')
    d = Document(file_path, text_params)
    txt = d.parse()
    family_name = d.find_fio()
    # new_txt = str.upper(txt)
    new_txt = open(os.path.join(*cwd, 'res.txt'), encoding='utf-8').read()
    d.change_text(new_txt)
    d.save(os.path.join(*cwd, 'my_doc.docx'))
    # print(txt)
    # print(d.find_fio())
